#!/usr/bin/env python3
"""
Generate a MiniCheck Analysis Report by copying relevant sections
verbatim from the TPO template document (preserving all XML formatting).

Critical findings (C=X) are passed in at runtime — the script discovers
which TPO Heading 2 sections to include by scanning each section's body
text for Mxxxx CHID references.  No hardcoded CHID list is needed.

Usage:
    python generate_minicheck_report.py [output_path] [chid,chid,...]

    output_path   Path for the generated .docx (default: MiniCheck_Report.docx)
    chid list     Comma-separated critical CHIDs, e.g. M0012,M0110,M0910
                  If omitted the script reads minicheck_results.json from the
                  project directory (written by the skill before calling this).
"""

import copy
import json
import os
import re
import sys
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ──────────────────────────────────────────────────────────────────────────────
# Configuration
# ──────────────────────────────────────────────────────────────────────────────
PROJECT_DIR   = '/Users/I321356/Documents/claude_project'
TEMPLATE_PATH = os.path.join(PROJECT_DIR, 'TPO_template.docx')
OUTPUT_PATH   = os.path.join(PROJECT_DIR, 'MiniCheck_Report.docx')
RESULTS_PATH  = os.path.join(PROJECT_DIR, 'minicheck_results.json')

COVER_INFO = {
    'title':  'SAP HANA MiniCheck Analysis Report',
    'system': 'HAN',
    'date':   '2026/04/06',
    'author': 'Quincy Chen',
}

HEADING_STYLES = {'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4',
                  'berschrift 1', 'berschrift 2', 'berschrift 3', 'berschrift 4'}

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

CHID_RE = re.compile(r'\bM\d{4}\b')


# ──────────────────────────────────────────────────────────────────────────────
# TPO template index: build chid → [section_title, ...] map at runtime
# ──────────────────────────────────────────────────────────────────────────────

def build_tpo_index(doc):
    """
    Scan every H2 section in the TPO template.
    CHIDs are extracted ONLY from the 'Check IDs:' line in each section body
    (not from free-form body text) to ensure accurate matching.
    Return a dict:  chid (str) -> list of section titles that list it.
    Also return:    section_title -> list of chids listed in that section.
    """
    chid_to_sections = {}   # M0012 -> ['Outdated SAP HANA Patch Level', ...]
    section_to_chids = {}   # 'Outdated SAP HANA Patch Level' -> ['M0012']

    paras = doc.paragraphs
    i = 0
    while i < len(paras):
        p = paras[i]
        if _is_heading(p) and _heading_level(p) == 2:
            title = _para_text(p)
            chids_found = []
            j = i + 1
            while j < len(paras):
                p2 = paras[j]
                if _is_heading(p2) and _heading_level(p2) <= 2:
                    break
                text = _para_text(p2)
                # Only extract CHIDs from the dedicated "Check IDs:" line
                if re.match(r'(?i)check\s+ids?\s*:', text):
                    chids_found = list(dict.fromkeys(CHID_RE.findall(text)))
                    break  # only one Check IDs line per section
                j += 1
            section_to_chids[title] = chids_found
            for chid in chids_found:
                chid_to_sections.setdefault(chid, []).append(title)
        i += 1

    return chid_to_sections, section_to_chids


def resolve_sections_for_critical_chids(critical_chids, chid_to_sections):
    """
    For each critical CHID, find the best matching TPO section title.
    Returns:
        ordered list of (chid, section_title) — one entry per CHID,
            section_title is None if no TPO section references this CHID.
        unique_sections: ordered list of section titles (deduped, CHID order)
    """
    mapping = []          # (chid, section_title or None)
    seen_sections = {}    # section_title -> first chid that claimed it

    for chid in critical_chids:
        sections = chid_to_sections.get(chid, [])
        if sections:
            # Prefer the first section that mentions this CHID directly
            section = sections[0]
            mapping.append((chid, section))
            seen_sections.setdefault(section, chid)
        else:
            mapping.append((chid, None))

    # Build unique ordered section list
    unique_sections = list(dict.fromkeys(s for _, s in mapping if s))

    return mapping, unique_sections


# ──────────────────────────────────────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────────────────────────────────────

def _is_heading(para):
    return para.style.name in HEADING_STYLES or para.style.name.startswith('Heading')


def _heading_level(para):
    m = re.search(r'\d+', para.style.name)
    return int(m.group()) if m else 99


def _para_text(para):
    return ''.join(r.text or '' for r in para.runs).strip()


def find_exact_section(doc, exact_title):
    """
    Find the H2 paragraph whose text exactly matches exact_title,
    then collect all body elements up to (but not including) the next H2 or H1.
    Returns list of lxml elements, or None if not found.
    """
    body_children = list(doc.element.body)
    elem_to_para = {id(p._element): p for p in doc.paragraphs}

    for i, elem in enumerate(body_children):
        if id(elem) not in elem_to_para:
            continue
        para = elem_to_para[id(elem)]
        if not _is_heading(para) or _heading_level(para) != 2:
            continue
        if _para_text(para) != exact_title:
            continue

        section_elems = []
        for j in range(i, len(body_children)):
            child = body_children[j]
            if j > i and id(child) in elem_to_para:
                p2 = elem_to_para[id(child)]
                if _is_heading(p2) and _heading_level(p2) <= 2:
                    break
            section_elems.append(child)
        return section_elems

    return None


def find_heading1_section(doc, h1_text):
    """Find an H1 containing h1_text and collect all elements until the next H1."""
    body_children = list(doc.element.body)
    elem_to_para = {id(p._element): p for p in doc.paragraphs}

    for i, elem in enumerate(body_children):
        if id(elem) not in elem_to_para:
            continue
        para = elem_to_para[id(elem)]
        if not _is_heading(para) or _heading_level(para) != 1:
            continue
        if h1_text.lower() not in _para_text(para).lower():
            continue

        section_elems = []
        for j in range(i, len(body_children)):
            child = body_children[j]
            if j > i and id(child) in elem_to_para:
                p2 = elem_to_para[id(child)]
                if _is_heading(p2) and _heading_level(p2) == 1:
                    break
            section_elems.append(child)
        return section_elems

    return None


def find_cover_elements(doc):
    """Return all body elements before the first H1 or ToC heading."""
    body_children = list(doc.element.body)
    elem_to_para = {id(p._element): p for p in doc.paragraphs}
    result = []
    for elem in body_children:
        if id(elem) in elem_to_para:
            para = elem_to_para[id(elem)]
            if _is_heading(para) and _heading_level(para) == 1:
                break
            t = _para_text(para).strip().lower()
            if t in ('table of contents', 'contents', 'inhaltsverzeichnis'):
                break
        result.append(elem)
    return result


def find_toc_heading(doc):
    for para in doc.paragraphs:
        t = _para_text(para).strip().lower()
        if t in ('table of contents', 'contents', 'inhaltsverzeichnis', 'table of content'):
            return copy.deepcopy(para._element)
    return None


def append_elements(target_body, elements):
    for elem in elements:
        target_body.append(copy.deepcopy(elem))


# ──────────────────────────────────────────────────────────────────────────────
# Word auto-ToC field
# ──────────────────────────────────────────────────────────────────────────────

def _make_toc_field():
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), 'TOCHeading')
    pPr.append(pStyle)
    p.append(pPr)

    r1 = OxmlElement('w:r')
    fc1 = OxmlElement('w:fldChar')
    fc1.set(qn('w:fldCharType'), 'begin')
    fc1.set(qn('w:dirty'), 'true')
    r1.append(fc1)
    p.append(r1)

    r2 = OxmlElement('w:r')
    instr = OxmlElement('w:instrText')
    instr.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    r2.append(instr)
    p.append(r2)

    r3 = OxmlElement('w:r')
    fc3 = OxmlElement('w:fldChar')
    fc3.set(qn('w:fldCharType'), 'separate')
    r3.append(fc3)
    p.append(r3)

    r4 = OxmlElement('w:r')
    rPr4 = OxmlElement('w:rPr')
    rPr4.append(OxmlElement('w:noProof'))
    r4.append(rPr4)
    t4 = OxmlElement('w:t')
    t4.text = 'Right-click here and choose "Update Field" to refresh table of contents'
    r4.append(t4)
    p.append(r4)

    r5 = OxmlElement('w:r')
    fc5 = OxmlElement('w:fldChar')
    fc5.set(qn('w:fldCharType'), 'end')
    r5.append(fc5)
    p.append(r5)

    return p


# ──────────────────────────────────────────────────────────────────────────────
# Action Plan table builder
# ──────────────────────────────────────────────────────────────────────────────

def _find_action_plan_table(doc):
    body_children = list(doc.element.body)
    elem_to_para = {id(p._element): p for p in doc.paragraphs}
    tables_map = {id(t._element): t for t in doc.tables}

    found_heading = False
    heading_idx = -1

    for i, elem in enumerate(body_children):
        if not found_heading:
            if id(elem) in elem_to_para:
                para = elem_to_para[id(elem)]
                if _is_heading(para) and 'action plan' in _para_text(para).lower():
                    found_heading = True
                    heading_idx = i
        else:
            if id(elem) in tables_map:
                tbl = tables_map[id(elem)]
                if len(tbl.columns) >= 6:
                    return tbl, heading_idx, i
    return None, -1, -1


def _set_cell_text(tc_elem, text):
    paras = tc_elem.findall(f'{{{W_NS}}}p')
    if not paras:
        p = OxmlElement('w:p')
        tc_elem.append(p)
        paras = [p]
    p = paras[0]
    for r in list(p.findall(f'{{{W_NS}}}r')):
        p.remove(r)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    for extra_p in paras[1:]:
        tc_elem.remove(extra_p)


def _make_action_plan_section(src_doc, unique_sections, sid):
    """
    Rebuild the Action Plan table keeping only rows matching the unique
    section titles derived from critical CHIDs.
    """
    table, heading_idx, table_idx = _find_action_plan_table(src_doc)
    if table is None:
        print('  [WARNING] Action Plan table not found in template')
        return []

    body_children = list(src_doc.element.body)
    result_elems = []

    for i in range(heading_idx, table_idx):
        result_elems.append(copy.deepcopy(body_children[i]))

    src_rows = table.rows
    if not src_rows:
        return result_elems

    # Build lookup: normalised issue title -> template row element
    template_rows_by_issue = {}
    for src_row in src_rows[1:]:
        issue_text = src_row.cells[4].text.strip()
        if issue_text and issue_text not in template_rows_by_issue:
            template_rows_by_issue[issue_text.lower()] = src_row._element

    fallback_row = src_rows[1]._element if len(src_rows) > 1 else src_rows[0]._element

    new_table = copy.deepcopy(table._element)
    tbl_elem = new_table
    for tr in list(tbl_elem.findall(f'{{{W_NS}}}tr')):
        tbl_elem.remove(tr)

    header_tr = copy.deepcopy(src_rows[0]._element)
    tbl_elem.append(header_tr)

    for seq_num, heading in enumerate(unique_sections, start=1):
        src_tr = template_rows_by_issue.get(heading.lower())
        if src_tr is None:
            # partial match fallback
            for key, elem in template_rows_by_issue.items():
                if heading.lower() in key or key in heading.lower():
                    src_tr = elem
                    break

        if src_tr is not None:
            new_tr = copy.deepcopy(src_tr)
            cells = new_tr.findall(f'.//{{{W_NS}}}tc')
            if len(cells) >= 6:
                _set_cell_text(cells[0], str(seq_num))
                db_val = ''.join(
                    n.text or '' for n in cells[1].iter(f'{{{W_NS}}}t')
                ).strip()
                _set_cell_text(cells[1], sid if db_val in ('-', '') else db_val)
        else:
            new_tr = copy.deepcopy(fallback_row)
            cells = new_tr.findall(f'.//{{{W_NS}}}tc')
            if len(cells) >= 6:
                _set_cell_text(cells[0], str(seq_num))
                _set_cell_text(cells[1], sid)
                _set_cell_text(cells[2], '-')
                _set_cell_text(cells[3], 'O')
                _set_cell_text(cells[4], heading)
                _set_cell_text(cells[5], '-')

        # Enforce left alignment on Issue and Recommendation columns
        cells = new_tr.findall(f'.//{{{W_NS}}}tc')
        for col_idx in (4, 5):
            if col_idx < len(cells):
                for p_elem in cells[col_idx].findall(f'.//{{{W_NS}}}p'):
                    pPr = p_elem.find(f'{{{W_NS}}}pPr')
                    if pPr is None:
                        pPr = OxmlElement('w:pPr')
                        p_elem.insert(0, pPr)
                    jc = pPr.find(f'{{{W_NS}}}jc')
                    if jc is None:
                        jc = OxmlElement('w:jc')
                        pPr.append(jc)
                    jc.set(f'{{{W_NS}}}val', 'left')
        tbl_elem.append(new_tr)

    result_elems.append(tbl_elem)
    return result_elems


def add_placeholder_section(target_body, title, chid):
    h_para = OxmlElement('w:p')
    h_pPr = OxmlElement('w:pPr')
    h_style = OxmlElement('w:pStyle')
    h_style.set(qn('w:val'), 'Heading2')
    h_pPr.append(h_style)
    h_para.append(h_pPr)
    h_run = OxmlElement('w:r')
    h_t = OxmlElement('w:t')
    h_t.text = title
    h_run.append(h_t)
    h_para.append(h_run)
    target_body.append(h_para)

    n_para = OxmlElement('w:p')
    n_run = OxmlElement('w:r')
    n_t = OxmlElement('w:t')
    n_t.text = f'[MiniCheck {chid}] — no matching section found in TPO template.'
    n_run.append(n_t)
    n_para.append(n_run)
    target_body.append(n_para)
    target_body.append(OxmlElement('w:p'))


# ──────────────────────────────────────────────────────────────────────────────
# Main
# ──────────────────────────────────────────────────────────────────────────────

def main(output_path=None, critical_chids=None):
    global OUTPUT_PATH
    if output_path:
        if not os.path.isabs(output_path):
            output_path = os.path.join(PROJECT_DIR, output_path)
        OUTPUT_PATH = output_path

    # ── Load critical CHIDs ───────────────────────────────────────────────────
    if critical_chids is None:
        if os.path.exists(RESULTS_PATH):
            with open(RESULTS_PATH) as f:
                results = json.load(f)
            critical_chids = sorted(
                {r['CHID'] for r in results if r.get('C', '').strip() == 'X' and r.get('CHID', '').strip()},
                key=lambda c: int(c[1:]) if c[1:].isdigit() else 0
            )
            print(f'Loaded {len(critical_chids)} critical CHIDs from {RESULTS_PATH}')
        else:
            print(f'[ERROR] No critical CHIDs provided and {RESULTS_PATH} not found.')
            sys.exit(1)

    print(f'Critical CHIDs ({len(critical_chids)}): {", ".join(critical_chids)}')

    # ── Open template and build TPO index ─────────────────────────────────────
    print(f'\nOpening template: {TEMPLATE_PATH}')
    src_doc = Document(TEMPLATE_PATH)

    print('Building TPO section index…')
    chid_to_sections, section_to_chids = build_tpo_index(src_doc)

    # ── Resolve critical CHIDs to TPO sections ────────────────────────────────
    chid_mapping, unique_sections = resolve_sections_for_critical_chids(
        critical_chids, chid_to_sections
    )

    print(f'\nSection resolution:')
    no_section = []
    for chid, section in chid_mapping:
        if section:
            print(f'  {chid} -> {section!r}')
        else:
            print(f'  {chid} -> [NO MATCHING SECTION IN TPO TEMPLATE]')
            no_section.append(chid)

    print(f'\n{len(unique_sections)} unique sections to include.')
    if no_section:
        print(f'{len(no_section)} CHIDs with no TPO section: {", ".join(no_section)}')

    # ── Build output document ─────────────────────────────────────────────────
    print('\nCreating output document from template (styles/themes preserved)…')
    out_doc = Document(TEMPLATE_PATH)
    body = out_doc.element.body
    for child in list(body):
        body.remove(child)

    sid = COVER_INFO.get('system', '-')

    # ── 1. Cover page ─────────────────────────────────────────────────────────
    print('\n[1/6] Copying cover page…')
    cover_elems = find_cover_elements(src_doc)
    print(f'      {len(cover_elems)} elements copied')
    append_elements(body, cover_elems)

    # ── 2. Table of Contents ──────────────────────────────────────────────────
    print('\n[2/6] Adding Table of Contents heading + auto-field…')
    toc_heading = find_toc_heading(src_doc)
    if toc_heading is not None:
        body.append(toc_heading)
    else:
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), 'Heading1')
        pPr.append(pStyle)
        p.append(pPr)
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = 'Table of Contents'
        r.append(t)
        p.append(r)
        body.append(p)

    body.append(_make_toc_field())
    body.append(OxmlElement('w:p'))

    # ── 3. Service Summary ────────────────────────────────────────────────────
    print('\n[3/6] Copying Service Summary section…')
    exec_summary = find_heading1_section(src_doc, 'Service Summary')
    if exec_summary:
        elem_to_para = {id(p._element): p for p in src_doc.paragraphs}
        ap_heading_local = None
        for i, elem in enumerate(exec_summary):
            if id(elem) in elem_to_para:
                p = elem_to_para[id(elem)]
                if _is_heading(p) and 'action plan' in _para_text(p).lower():
                    ap_heading_local = i
                    break

        if ap_heading_local is not None:
            append_elements(body, exec_summary[:ap_heading_local])
            print('      Rebuilding Action Plan table…')
            ap_elems = _make_action_plan_section(src_doc, unique_sections, sid)
            for e in ap_elems:
                body.append(e)
        else:
            append_elements(body, exec_summary)
        print(f'      Service Summary: {len(exec_summary)} elements')
    else:
        print('      [WARNING] Service Summary not found in template')

    # ── 4. General Overview ───────────────────────────────────────────────────
    print('\n[4/6] Copying General Overview section…')
    gen_overview = find_heading1_section(src_doc, 'General Overview')
    if gen_overview:
        append_elements(body, gen_overview)
        print(f'      {len(gen_overview)} elements copied')
    else:
        print('      [WARNING] General Overview not found in template')

    # ── 5. Check Lists ────────────────────────────────────────────────────────
    print('\n[5/6] Copying Check Lists section…')
    check_lists = find_heading1_section(src_doc, 'Check List')
    if check_lists:
        append_elements(body, check_lists)
        print(f'      {len(check_lists)} elements copied')
    else:
        print('      [WARNING] Check Lists not found in template')

    # ── 6. Issues and Recommendations ────────────────────────────────────────
    print('\n[6/6] Adding Issues and Recommendations sections…')
    issues_h1 = find_heading1_section(src_doc, 'Issues and Recommendation')
    if issues_h1:
        body.append(copy.deepcopy(issues_h1[0]))
    else:
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), 'Heading1')
        pPr.append(pStyle)
        p.append(pPr)
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = 'Issues and Recommendations'
        r.append(t)
        p.append(r)
        body.append(p)

    not_found = []
    added_headings = set()

    for section_title in unique_sections:
        if section_title in added_headings:
            continue
        elems = find_exact_section(src_doc, section_title)
        # Identify which CHIDs map to this section (for logging)
        chids_here = [c for c, s in chid_mapping if s == section_title]
        if elems:
            append_elements(body, elems)
            added_headings.add(section_title)
            print(f'      [OK] {", ".join(chids_here)}: {section_title!r} ({len(elems)} elements)')
        else:
            not_found.append((chids_here, section_title))
            print(f'      [--] {", ".join(chids_here)}: {section_title!r} NOT FOUND — adding placeholder')
            add_placeholder_section(body, section_title, '/'.join(chids_here))

    # CHIDs with no section at all — add placeholder
    for chid in no_section:
        desc = chid  # just use CHID as title since we have no section
        print(f'      [--] {chid}: no TPO section found — adding placeholder')
        add_placeholder_section(body, f'[{chid}] No TPO section available', chid)

    # ── Save ──────────────────────────────────────────────────────────────────
    print(f'\nSaving output to: {OUTPUT_PATH}')
    out_doc.save(OUTPUT_PATH)
    print('Done — output file created successfully.')

    return not_found


if __name__ == '__main__':
    args = sys.argv[1:]
    out = None
    chids = None

    for arg in args:
        if ',' in arg or (arg.startswith('M') and len(arg) == 5):
            # Looks like a CHID list
            chids = [c.strip() for c in arg.split(',') if c.strip()]
        else:
            out = arg

    not_found = main(output_path=out, critical_chids=chids)
    if not_found:
        print('\nSections not found in template:')
        for chids_nf, title in not_found:
            print(f'  {", ".join(chids_nf)}: {title}')
